import errno
import logging
import re
import os
import shutil
import subprocess
from typing import NamedTuple

import xlsxwriter
from docx import Document
from difflib import SequenceMatcher

class DocumentRow(NamedTuple):
    type: list
    text: list
    justification: list

class DocumentLink(NamedTuple):
    link: str
    channel: str
    justification: list

DocumentRows = []
DocumentLinks = []

logger = logging.getLogger(__name__)

# Словарь для замены одинаково выглядящих кириллических букв на латинские
cyrillic_to_latin_simple = {
    'А': 'A', 'В': 'B', 'Е': 'E', 'К': 'K', 'М': 'M', 'Н': 'H', 'О': 'O', 'Р': 'P', 'С': 'C', 'Т': 'T', 'У': 'Y', 'Х': 'X',
    'а': 'a', 'в': 'b', 'е': 'e', 'к': 'k', 'м': 'm', 'н': 'h', 'о': 'o', 'р': 'p', 'с': 'c', 'т': 't', 'у': 'y', 'х': 'x'
}

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

def is_valid_number(text, start, end):
    before = text[start - 1] if start > 0 else ' '
    after = text[end] if end < len(text) else ' '
    return before in ' \t\n.,;:!?()[]{}' and after in ' \t\n.,;:!?()[]{}'

def remove_point(link):
    if len(link) == 0:
        return link
    if link[-1] == '.':
        link = link[:-1]
    return link

def extract_date(justification):
    date_pattern = re.compile(
        r'(\d{1,2}(?:\s*\.\s*|\s+)(?:января|февраля|марта|апреля|мая|июня|июля|августа|сентября|октября|ноября|декабря|\d{1,2})(?:\s*\.\s*|\s+)\d{4})',
        re.IGNORECASE
    )
    
    match = date_pattern.search(justification)
    if match:
        return "Решение от " + match.group() + " г."
    
    return justification

def remove_disclaimer(justification):
    disclaimer_index = justification.find('Подлежит немедленному исполнению')
    if disclaimer_index != -1:
        justification = justification[:disclaimer_index]
    
    justification = extract_date(justification)
    
    return justification

def replace_cyrillic_with_latin(link):
    return ''.join(cyrillic_to_latin_simple.get(char, char) for char in link)

def load_document(filename):
    logger.log(level=logging.INFO, msg="Начало загрузки документа")

    DocumentRows.clear()
    DocumentLinks.clear()

    document = Document(filename)
    rows = 0
    
    link_pattern = re.compile(r'(https?://)?t\.me/(s/)?([a-zA-Zа-яА-Я0-9_.-]+)')
    id_pattern = re.compile(r'\d{5,}')

    for table in document.tables:
        for row in table.rows:
            rows += 1
            row_cells = []
            for cell in row.cells:
                cell_para = []
                for para in cell.paragraphs:
                    cell_para.append(para.text)
                row_cells.append(cell_para)
            DocumentRows.append(DocumentRow(type=row_cells[0], text=row_cells[1], justification=row_cells[2]))

            for cell_para in row_cells[1]:
                links = []
                channel_id = -1
                channel_name = ''

                for match in link_pattern.finditer(cell_para):
                    match_str = match.group(0).strip()  # Убираем пробелы по краям

                    # Проверяем наличие ссылки с /s/
                    if re.search(r'(https?://)?t\.me/s/', match_str):
                        print(f"Найдена ссылка с /s/: {match_str}")

                        # Приводим к стандартному виду и удаляем /s/
                        match_str = re.sub(r'(https?://)?t\.me/s/', 'https://t.me/', match_str)
                        print(f"Ссылка после замены: {match_str}")

                        # Применяем регулярное выражение повторно к измененной строке
                        new_match = re.search(link_pattern, match_str)
                        if new_match:
                            channel_name = new_match.group(3)  # Берем название канала из измененной строки
                            print(f"Извлеченное название канала: {channel_name}")
                        else:
                            print(f"Ошибка: не удалось найти совпадение в обновленной строке {match_str}")
                    else:
                        channel_name = match.group(3)  # Если ссылка в правильном виде, берем третью группу сразу
                        print(f"Ссылка уже в правильном виде, название канала: {channel_name}")

                    full_link = f"https://t.me/{replace_cyrillic_with_latin(channel_name)}"
                    print(f"Форматированная ссылка: {full_link}")
                    links.append(remove_point(full_link))  # Убираем точку в конце, если она есть
                    print(f"Ссылка добавлена в список: {full_link}\n")

                    for match in id_pattern.finditer(cell_para):
                        number = match.group()
                        if number.startswith("100"):
                            number = number[3:]
                        if is_valid_number(cell_para, match.start(), match.end()) and len(number) >= 5:
                            channel_id = number

                    if links:
                        for link in links:
                            DocumentLinks.append(
                                DocumentLink(
                                    link=link,
                                    channel=channel_id,
                                    justification=row_cells[2]
                                )
                            )

    logger.log(level=logging.INFO, msg="Конец загрузки документа, прочитано  " + str(rows) + " строк")
    print("Запись линков в файл ")
    write_links()

def is_not_blank(s):
    return bool(s and not s.isspace())

def silentremove(filename):
    try:
        os.remove(filename)
    except OSError as e:
        if e.errno != errno.ENOENT:
            raise

def write_links():
    silentremove('telegram-links.xlsx')
    workbook = xlsxwriter.Workbook('telegram-links.xlsx')
    worksheet = workbook.add_worksheet()

    worksheet.set_column('A:A', 100)
    worksheet.set_column('B:B', 50)
    worksheet.set_column('C:C', 200)

    worksheet.write('A1', 'Ссылки на каналы')
    worksheet.write('B1', 'ID канала')
    worksheet.write('C1', 'Решения суда')

    row = 1
    for link in DocumentLinks:
        worksheet.write(row, 0, link.link)
        if link.channel != '-1':
            worksheet.write(row, 1, link.channel)
        text = ''
        for para in link.justification:
            if is_not_blank(para):
                text += para + '\n'
        if is_not_blank(text):
            worksheet.write(row, 2, remove_disclaimer(text))
        row += 1

    workbook.close()
    shutil.copy2('telegram-links.xlsx', 'root/files_for_bots/telegram-links.xlsx')

def parse_document(text):
    foundRows = []
    logger.log(level=logging.INFO, msg="Начало поиска последовательности " + text)
    for row in DocumentRows:
        for para in row.text:
            if text.casefold() in para.casefold():
                foundRows.append(row)
                break

    logger.log(level=logging.INFO, msg="Конец поиска последовательности " + text)
    return foundRows

def get_row(n):
    return DocumentRows[n]
